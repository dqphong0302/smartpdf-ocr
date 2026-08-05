"""Smart PDF — FastAPI Backend."""
import asyncio
import logging
import os
import secrets
import shutil
import sqlite3
import time
import uuid
from collections import defaultdict, deque
from contextlib import asynccontextmanager
from pathlib import Path

from auth import authenticate, init_auth_tables, logout, seed_default_user, validate_session
from dotenv import load_dotenv
from fastapi import (
    BackgroundTasks,
    Depends,
    FastAPI,
    File,
    Form,
    HTTPException,
    Query,
    Request,
    UploadFile,
    WebSocket,
    WebSocketDisconnect,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from job_manager import Job, JobStatus, PageResult, PageStatus, job_manager
from latex_compiler import LatexCompileError, compile_latex_project, get_latex_health, prepare_latex_workspace
from ocr_engine import sanitize_ocr_html, smart_ocr, vision_ocr_batch
from pdf_analyzer import (
    analyze_pdf,
    extract_page_images,
    extract_page_markdown,
    extract_page_text,
    get_page_thumbnail,
    render_page_to_image,
)
from pydantic import BaseModel

load_dotenv()

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)
MAX_SIZE = int(os.getenv("MAX_UPLOAD_SIZE_MB", "50")) * 1024 * 1024
MAX_PDF_PAGES = int(os.getenv("MAX_PDF_PAGES", "500"))
JOB_MAX_AGE_DAYS = int(os.getenv("JOB_MAX_AGE_DAYS", "7"))
MAX_CONCURRENT_JOBS = max(1, int(os.getenv("MAX_CONCURRENT_JOBS", "2")))
ENABLE_API_DOCS = os.getenv("ENABLE_API_DOCS", "false").lower() in {"1", "true", "yes", "on"}
LATEX_COMPILE_ENABLED = os.getenv("LATEX_COMPILE_ENABLED", "false").lower() in {"1", "true", "yes", "on"}

logger = logging.getLogger("smart-pdf")


@asynccontextmanager
async def lifespan(_app: FastAPI):
    init_auth_tables()
    if not seed_default_user():
        logger.warning("Admin user was not seeded because explicit credentials are missing")
    cleanup_task = asyncio.create_task(_cleanup_scheduler())
    try:
        yield
    finally:
        cleanup_task.cancel()
        try:
            await cleanup_task
        except asyncio.CancelledError:
            pass


app = FastAPI(
    title="Smart PDF",
    version="1.1.0",
    docs_url="/docs" if ENABLE_API_DOCS else None,
    redoc_url="/redoc" if ENABLE_API_DOCS else None,
    openapi_url="/openapi.json" if ENABLE_API_DOCS else None,
    lifespan=lifespan,
)

COOKIE_SECURE = os.getenv("COOKIE_SECURE", "true").lower() in {"1", "true", "yes", "on"}
COOKIE_SAMESITE = os.getenv("COOKIE_SAMESITE", "lax")
CORS_ORIGINS = [
    origin.strip()
    for origin in os.getenv("CORS_ORIGINS", "*").split(",")
    if origin.strip()
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS or ["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

_job_semaphore = asyncio.Semaphore(MAX_CONCURRENT_JOBS)
_login_attempts: dict[str, deque[float]] = defaultdict(deque)
LOGIN_WINDOW_SECONDS = 15 * 60
LOGIN_MAX_ATTEMPTS = 5


@app.middleware("http")
async def security_headers(request: Request, call_next):
    response = await call_next(request)
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault("X-Frame-Options", "DENY")
    response.headers.setdefault("Referrer-Policy", "no-referrer")
    response.headers.setdefault("Permissions-Policy", "camera=(), microphone=(), geolocation=()")
    response.headers.setdefault(
        "Content-Security-Policy",
        "default-src 'self'; img-src 'self' data: blob:; style-src 'self' 'unsafe-inline'; "
        "script-src 'self'; connect-src 'self' wss:; object-src 'none'; base-uri 'self'; frame-ancestors 'none'",
    )
    if COOKIE_SECURE:
        response.headers.setdefault("Strict-Transport-Security", "max-age=31536000; includeSubDomains")
    return response


async def read_upload_limited(file: UploadFile) -> bytes:
    """Read an upload in bounded chunks and reject oversized bodies early."""
    chunks = []
    total = 0
    while chunk := await file.read(1024 * 1024):
        total += len(chunk)
        if total > MAX_SIZE:
            raise HTTPException(413, f"File too large. Max {MAX_SIZE // (1024 * 1024)}MB")
        chunks.append(chunk)
    return b"".join(chunks)


def safe_upload_name(filename: str | None, allowed_suffixes: set[str]) -> str:
    """Return a normalized basename and enforce the declared file type."""
    safe_name = Path(filename or "").name.strip()
    if not safe_name or Path(safe_name).suffix.lower() not in allowed_suffixes:
        expected = ", ".join(sorted(allowed_suffixes))
        raise HTTPException(400, f"Only {expected} files are accepted")
    return safe_name


# ── Auth ────────────────────────────────────────────────────────────

async def _cleanup_scheduler():
    """Run cleanup every 24 hours to delete expired jobs."""
    from database import cleanup_expired_jobs
    # Run an initial cleanup on startup
    await asyncio.sleep(10)  # wait for app to fully start
    stats = cleanup_expired_jobs(JOB_MAX_AGE_DAYS)
    if stats["deleted_jobs"] > 0:
        logger.info(f"Startup cleanup: {stats['deleted_jobs']} jobs deleted, {stats['bytes_freed_mb']}MB freed")
    # Then run daily
    while True:
        await asyncio.sleep(86400)  # 24 hours
        try:
            stats = cleanup_expired_jobs(JOB_MAX_AGE_DAYS)
            if stats["deleted_jobs"] > 0:
                logger.info(f"Daily cleanup: {stats['deleted_jobs']} jobs deleted, {stats['bytes_freed_mb']}MB freed")
        except Exception as e:
            logger.error(f"Cleanup failed: {e}")


def require_auth(request: Request):
    """Dependency: require valid session cookie."""
    token = request.cookies.get("session")
    username = validate_session(token)
    if not username:
        raise HTTPException(401, "Chưa đăng nhập")
    return username


class LoginBody(BaseModel):
    username: str
    password: str


@app.post("/api/auth/login")
async def login(body: LoginBody, request: Request):
    client_key = request.client.host if request.client else "unknown"
    now = time.time()
    attempts = _login_attempts[client_key]
    while attempts and attempts[0] < now - LOGIN_WINDOW_SECONDS:
        attempts.popleft()
    if len(attempts) >= LOGIN_MAX_ATTEMPTS:
        raise HTTPException(429, "Too many login attempts. Try again later.")

    token = authenticate(body.username, body.password)
    if not token:
        attempts.append(now)
        raise HTTPException(401, "Sai tên đăng nhập hoặc mật khẩu")
    attempts.clear()
    resp = JSONResponse({"status": "ok", "username": body.username})
    resp.set_cookie(
        "session", token,
        httponly=True,
        max_age=30 * 24 * 3600,
        samesite=COOKIE_SAMESITE,
        secure=COOKIE_SECURE,
        path="/",
    )
    return resp


@app.post("/api/auth/logout")
async def logout_endpoint(request: Request):
    token = request.cookies.get("session")
    if token:
        logout(token)
    resp = JSONResponse({"status": "ok"})
    resp.delete_cookie("session", path="/", secure=COOKIE_SECURE, samesite=COOKIE_SAMESITE)
    return resp


@app.get("/api/auth/check")
async def check_auth(request: Request):
    token = request.cookies.get("session")
    username = validate_session(token)
    if not username:
        raise HTTPException(401, "Chưa đăng nhập")
    return {"status": "ok", "username": username}


# ── Upload PDF ──────────────────────────────────────────────────────
@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...), _user: str = Depends(require_auth)):
    filename = safe_upload_name(file.filename, {".pdf"})
    content = await read_upload_limited(file)

    # Save file
    filepath = UPLOAD_DIR / f"{uuid.uuid4().hex[:12]}_{filename}"
    with open(filepath, "wb") as f:
        f.write(content)

    # Create job
    job = job_manager.create_job(filename=filename, filepath=str(filepath))

    # Analyze PDF
    await job_manager.update_job_status(job.job_id, JobStatus.ANALYZING)
    try:
        analysis = analyze_pdf(str(filepath))
        if analysis.total_pages > MAX_PDF_PAGES:
            raise ValueError(f"PDF has {analysis.total_pages} pages; maximum is {MAX_PDF_PAGES}")
    except Exception as e:
        await job_manager.update_job_status(job.job_id, JobStatus.FAILED)
        job.error = str(e)
        raise HTTPException(500, f"PDF analysis failed: {e}") from e

    job.total_pages = analysis.total_pages
    # Initialize page results
    for pa in analysis.pages:
        job.pages[pa.page_num] = PageResult(
            page_num=pa.page_num,
            classification=pa.classification,
        )

    job.status = JobStatus.UPLOADED
    # Persist to SQLite
    job_manager.persist_job(job.job_id)
    for pa in analysis.pages:
        job_manager.persist_page(job.job_id, pa.page_num)

    return {
        "job_id": job.job_id,
        "filename": filename,
        "analysis": analysis.to_dict(),
    }


# ── List All Jobs ──────────────────────────────────────────────────
@app.get("/api/jobs")
async def list_jobs(_user: str = Depends(require_auth)):
    return job_manager.list_jobs()


# ── Get Job Status ─────────────────────────────────────────────────
@app.get("/api/jobs/{job_id}")
async def get_job(job_id: str, include_text: bool = False, _user: str = Depends(require_auth)):
    data = job_manager.get_job_dict(job_id, include_text=include_text)
    if not data:
        raise HTTPException(404, "Job not found")
    return data


# ── Delete Job ─────────────────────────────────────────────────────
@app.delete("/api/jobs/{job_id}")
async def delete_job(job_id: str, _user: str = Depends(require_auth)):
    if not job_manager.delete_job(job_id):
        raise HTTPException(404, "Job not found")
    return {"status": "deleted", "job_id": job_id}


# ── Start OCR Processing ───────────────────────────────────────────
@app.post("/api/ocr/{job_id}")
async def start_ocr(
    job_id: str,
    pages: list[int] = Query(default=None, description="Page numbers to process"),
    mode: str = Query(default="all", description="all|odd|even|custom"),
    _user: str = Depends(require_auth),
    force_method: str = Query(default=None, description="tesseract|vision|auto"),
    extract_images: bool = Query(default=False, description="Extract original images"),
):
    job = job_manager.get_job(job_id)
    if not job:
        raise HTTPException(404, "Job not found")
    if job.status == JobStatus.PROCESSING:
        raise HTTPException(409, "Job is already processing")

    # Determine which pages to process
    all_pages = list(range(1, job.total_pages + 1))

    if mode == "odd":
        selected = [p for p in all_pages if p % 2 == 1]
    elif mode == "even":
        selected = [p for p in all_pages if p % 2 == 0]
    elif mode == "custom" and pages:
        selected = [p for p in pages if 1 <= p <= job.total_pages]
    else:
        selected = all_pages

    job.selected_pages = selected

    # Mark non-selected pages as skipped and persist to DB
    for p_num in all_pages:
        if p_num not in selected:
            job.pages[p_num].status = PageStatus.COMPLETED
            job.pages[p_num].method = "skipped"
            job_manager.persist_page(job.job_id, p_num)

    # Persist job state (selected_pages) to DB
    job_manager.persist_job(job.job_id)

    # Process in background
    asyncio.create_task(_process_pages_limited(job, selected, force_method, extract_images))

    return {"job_id": job_id, "selected_pages": selected, "total": len(selected)}


PARALLEL_BATCHES = int(os.getenv("PARALLEL_BATCHES", "4"))  # concurrent vision API requests
BATCH_SIZE = int(os.getenv("BATCH_SIZE", "4"))  # pages per vision API call


def inline_base64_images(html_content: str, job_id: str) -> str:
    """Find all references to /api/extracted-images/{job_id}/{filename} in html_content,
    read them from disk, base64-encode them, and inline them into the src attribute."""
    import base64
    import os
    import re

    # Pattern matches src="/api/extracted-images/{job_id}/{filename}"
    pattern = r'src=["\']/api/extracted-images/([a-zA-Z0-9_\-]+)/([^"\']+)["\']'

    def replace_img(match):
        matched_job_id = match.group(1)
        filename = os.path.basename(match.group(2))
        img_path = UPLOAD_DIR / "extracted_images" / matched_job_id / filename
        if img_path.exists() and img_path.is_file():
            try:
                suffix = img_path.suffix.lstrip(".").lower()
                if suffix in ("jpg", "jpeg"):
                    mime = "image/jpeg"
                elif suffix == "png":
                    mime = "image/png"
                elif suffix == "gif":
                    mime = "image/gif"
                elif suffix == "webp":
                    mime = "image/webp"
                else:
                    mime = f"image/{suffix}"

                with open(img_path, "rb") as f:
                    b64_data = base64.b64encode(f.read()).decode("utf-8")
                return f'src="data:{mime};base64,{b64_data}"'
            except Exception as e:
                logger.error(f"Error encoding extracted image to base64 {img_path}: {e}")
        return match.group(0)

    return re.sub(pattern, replace_img, html_content)


def _enrich_html_with_images(job_id: str, page_num: int, filepath: str, html_text: str, extract_images: bool) -> str:
    """Helper to extract page images using PyMuPDF and append a styled gallery to the page HTML."""
    html_text = sanitize_ocr_html(html_text)
    if not extract_images:
        return html_text
    
    out_dir = UPLOAD_DIR / "extracted_images" / job_id
    try:
        saved_files = extract_page_images(filepath, page_num, str(out_dir))
    except Exception as e:
        logger.error(f"Error extracting images for job {job_id} page {page_num}: {e}")
        saved_files = []
        
    if not saved_files:
        return html_text
        
    # Generate HTML gallery
    gallery_html = '<div class="extracted-images-gallery" data-page="' + str(page_num) + '">'
    for fname in saved_files:
        img_url = f"/api/extracted-images/{job_id}/{fname}"
        gallery_html += (
            f'<div class="extracted-image-item" onclick="window.openLightbox(\'{img_url}\')">'
            f'<img src="{img_url}" alt="Page {page_num} Extracted Image" class="extracted-img" />'
            f'</div>'
        )
    gallery_html += '</div>'
    return f"{html_text}\n{gallery_html}"


async def _process_pages(job: Job, pages: list[int], force_method: str = None, extract_images: bool = False):
    """Background task: process selected pages with smart OCR and optional image extraction.
    
    Phase 1: Process digital/tesseract pages sequentially (fast, no API).
    Phase 2: Collect all vision pages, split into batches of BATCH_SIZE,
             then run up to PARALLEL_BATCHES concurrently via asyncio.gather.
    """
    import time as _time
    await job_manager.update_job_status(job.job_id, JobStatus.PROCESSING)

    # ── Phase 1: Handle digital & tesseract pages (fast, sequential) ──
    vision_pages = []  # collect (page_num, image) for batched vision processing
    
    for page_num in pages:
        page_result = job.pages[page_num]
        classification = page_result.classification

        await job_manager.update_page(
            job.job_id, page_num, status=PageStatus.PROCESSING
        )

        try:
            # Digital pages: extract text directly (no OCR needed)
            if classification == "digital" and force_method != "vision":
                import mistune
                start = _time.time()

                # Try pdf-inspector Markdown extraction first (layout-aware: tables, columns)
                md = extract_page_markdown(job.filepath, page_num)
                if md:
                    text = md  # store Markdown as the text representation
                    raw_html = mistune.html(md)
                else:
                    # Fallback: plain PyMuPDF text wrapped in <pre>
                    text = extract_page_text(job.filepath, page_num)
                    raw_html = f"<pre>{text}</pre>"

                elapsed = _time.time() - start
                enriched_html = _enrich_html_with_images(job.job_id, page_num, job.filepath, raw_html, extract_images)

                await job_manager.update_page(
                    job.job_id,
                    page_num,
                    status=PageStatus.COMPLETED,
                    method="digital",
                    text=text,
                    html_text=enriched_html,
                    confidence=100.0,
                    time_taken=elapsed,
                )
                continue

            # If force_method is tesseract, process individually
            if force_method == "tesseract":
                image = render_page_to_image(job.filepath, page_num)
                result = await smart_ocr(image, classification, force_method="tesseract")
                
                raw_html = result.get("html_text", f"<pre>{result['text']}</pre>")
                enriched_html = _enrich_html_with_images(job.job_id, page_num, job.filepath, raw_html, extract_images)
                
                await job_manager.update_page(
                    job.job_id,
                    page_num,
                    status=PageStatus.COMPLETED,
                    method=result["method"],
                    text=result["text"],
                    html_text=enriched_html,
                    confidence=result["confidence"],
                    time_taken=result["time_taken"],
                )
                continue

            # For auto/vision: try tesseract first on simple scans
            if classification == "scan_simple" and force_method != "vision":
                image = render_page_to_image(job.filepath, page_num)
                result = await smart_ocr(image, classification, force_method=None)
                if result["method"] == "tesseract":
                    # Tesseract was good enough
                    raw_html = result.get("html_text", f"<pre>{result['text']}</pre>")
                    enriched_html = _enrich_html_with_images(job.job_id, page_num, job.filepath, raw_html, extract_images)
                    
                    await job_manager.update_page(
                        job.job_id,
                        page_num,
                        status=PageStatus.COMPLETED,
                        method=result["method"],
                        text=result["text"],
                        html_text=enriched_html,
                        confidence=result["confidence"],
                        time_taken=result["time_taken"],
                    )
                    continue
                # If smart_ocr fell back to vision, it already did the API call
                raw_html = result.get("html_text", f"<pre>{result['text']}</pre>")
                enriched_html = _enrich_html_with_images(job.job_id, page_num, job.filepath, raw_html, extract_images)
                
                await job_manager.update_page(
                    job.job_id,
                    page_num,
                    status=PageStatus.COMPLETED,
                    method=result["method"],
                    text=result["text"],
                    html_text=enriched_html,
                    confidence=result["confidence"],
                    time_taken=result["time_taken"],
                )
                continue

            # Vision-destined pages: render image and collect for parallel batching
            image = render_page_to_image(job.filepath, page_num)
            vision_pages.append((page_num, image))

        except Exception as e:
            await job_manager.update_page(
                job.job_id,
                page_num,
                status=PageStatus.FAILED,
                error=str(e),
            )

    # ── Phase 2: Parallel vision batch processing ──
    if vision_pages:
        # Split into batches of BATCH_SIZE pages each
        batches = [
            vision_pages[i:i + BATCH_SIZE]
            for i in range(0, len(vision_pages), BATCH_SIZE)
        ]

        logger.info(
            f"Job {job.job_id}: {len(vision_pages)} vision pages → "
            f"{len(batches)} batches (size={BATCH_SIZE}), "
            f"parallel={PARALLEL_BATCHES}"
        )

        async def _run_batch(batch):
            """Process a single vision batch and update page results."""
            try:
                results = await vision_ocr_batch(batch)
                for (pn, _img), result in zip(batch, results, strict=True):
                    raw_html = result.get("html_text", f"<pre>{result['text']}</pre>")
                    enriched_html = _enrich_html_with_images(job.job_id, pn, job.filepath, raw_html, extract_images)
                    
                    await job_manager.update_page(
                        job.job_id,
                        pn,
                        status=PageStatus.COMPLETED,
                        method=result["method"],
                        text=result["text"],
                        html_text=enriched_html,
                        confidence=result["confidence"],
                        time_taken=result["time_taken"],
                    )
            except Exception as e:
                for pn, _ in batch:
                    await job_manager.update_page(
                        job.job_id,
                        pn,
                        status=PageStatus.FAILED,
                        error=str(e),
                    )

        # Run batches in parallel waves of PARALLEL_BATCHES
        for wave_start in range(0, len(batches), PARALLEL_BATCHES):
            wave = batches[wave_start:wave_start + PARALLEL_BATCHES]
            await asyncio.gather(*[_run_batch(b) for b in wave])

    await job_manager.update_job_status(job.job_id, JobStatus.COMPLETED)


async def _process_pages_limited(job: Job, pages: list[int], force_method: str = None, extract_images: bool = False):
    async with _job_semaphore:
        await _process_pages(job, pages, force_method, extract_images)



# ── Page Thumbnail ──────────────────────────────────────────────────
@app.get("/api/thumbnail/{job_id}/{page_num}")
async def get_thumbnail(job_id: str, page_num: int, width: int = 200, _user: str = Depends(require_auth)):
    width = max(64, min(width, 1600))
    job = job_manager.get_job(job_id)
    filepath = job.filepath if job else None
    if not filepath:
        # Try DB
        data = job_manager.get_job_dict(job_id)
        if not data:
            raise HTTPException(404, "Job not found")
        # Need to get filepath from DB
        from database import _get_conn
        conn = _get_conn()
        row = conn.execute("SELECT filepath FROM jobs WHERE job_id = ?", (job_id,)).fetchone()
        conn.close()
        filepath = row["filepath"] if row else None
    if not filepath:
        raise HTTPException(404, "Job not found")

    png_bytes = get_page_thumbnail(filepath, page_num, width)
    return Response(content=png_bytes, media_type="image/png")


# ── Extracted Images Safe Serving ───────────────────────────────────
@app.get("/api/extracted-images/{job_id}/{filename}")
async def get_extracted_image(job_id: str, filename: str, _user: str = Depends(require_auth)):
    """Serve an extracted image safely from the uploads directory."""
    import os
    # Sandbox check: prevent directory traversal
    safe_filename = os.path.basename(filename)
    img_path = UPLOAD_DIR / "extracted_images" / job_id / safe_filename
    if not img_path.exists() or not img_path.is_file():
        raise HTTPException(404, "Extracted image not found")
    return FileResponse(str(img_path))


# ── Download Results ────────────────────────────────────────────────
@app.get("/api/download/{job_id}")
async def download_results(job_id: str, format: str = "txt", _user: str = Depends(require_auth)):
    import urllib.parse
    export_format = (format or "txt").lower()
    if export_format == "text":
        export_format = "txt"
    # Get full data including text
    data = job_manager.get_job_dict(job_id, include_text=True)
    if not data:
        # Fallback to API jobs (in-memory)
        data = _api_jobs.get(job_id)
        if not data:
            raise HTTPException(404, "Job not found")

    filename = data["filename"]
    pages = data["pages"]  # dict of str(page_num) -> page dict

    if export_format == "json":
        import json
        export = []
        for num in sorted(pages.keys(), key=int):
            p = pages[num]
            export.append({
                "page": int(num),
                "method": p.get("method"),
                "confidence": p.get("confidence", 0),
                "text": p.get("text", ""),
                "html_text": p.get("html_text", ""),
            })
        content = json.dumps(export, ensure_ascii=False, indent=2)
        filename_encoded = urllib.parse.quote(f"{filename}_ocr.json")
        return Response(
            content=content,
            media_type="application/json",
            headers={"Content-Disposition": f"attachment; filename*=utf-8''{filename_encoded}"},
        )
    elif export_format == "html":
        html_parts = [
            '<!DOCTYPE html>',
            '<html lang="vi">',
            '<head>',
            f'<title>OCR: {filename}</title>',
            '<meta charset="utf-8">',
            '<style>',
            'body { font-family: Georgia, serif; max-width: 900px; margin: 0 auto; padding: 20px; background: #fafafa; color: #333; }',
            '.page { background: white; padding: 30px; margin: 20px 0; border: 1px solid #ddd; box-shadow: 0 2px 8px rgba(0,0,0,0.08); }',
            '.page-header { border-bottom: 2px solid #eee; padding-bottom: 8px; margin-bottom: 16px; font-size: 14px; color: #888; }',
            'table { border-collapse: collapse; width: 100%; }',
            'td, th { border: 1px solid #ccc; padding: 6px 10px; }',
            '.extracted-images-gallery { display: grid; grid-template-columns: repeat(auto-fill, minmax(150px, 1fr)); gap: 12px; margin: 15px 0; padding: 10px; background: rgba(0, 0, 0, 0.02); border-radius: 8px; border: 1px solid rgba(0, 0, 0, 0.05); }',
            '.extracted-image-item { border-radius: 6px; overflow: hidden; border: 1px solid #eee; background: #fff; display: flex; align-items: center; justify-content: center; }',
            '.extracted-img { max-width: 100%; max-height: 200px; object-fit: contain; display: block; }',
            '</style>',
            '</head>',
            '<body>',
            f'<h1>📄 {filename}</h1>',
        ]
        for num in sorted(pages.keys(), key=int):
            p = pages[num]
            text = p.get("text", "")
            html_text = p.get("html_text", "")
            if text or html_text:
                html_parts.append('<div class="page">')
                html_parts.append(f'<div class="page-header">Trang {num} · {p.get("method", "")} · {p.get("confidence", 0)}%</div>')
                html_parts.append(html_text if html_text else f'<pre>{text}</pre>')
                html_parts.append('</div>')
        html_parts.append('</body></html>')
        content = '\n'.join(html_parts)
        # Inline images to Base64 so it is 100% self-contained
        content = inline_base64_images(content, job_id)
        filename_encoded = urllib.parse.quote(f"{filename}_ocr.html")
        return Response(
            content=content,
            media_type="text/html; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename*=utf-8''{filename_encoded}"},
        )
    elif export_format in {"markdown", "md"}:
        md_parts = [f"# {filename}\n"]
        for num in sorted(pages.keys(), key=int):
            p = pages[num]
            text = p.get("text", "")
            html_text = p.get("html_text", "")
            if text or html_text:
                md_parts.append(f"\n---\n\n## Trang {num} · {p.get('method', '')} · {p.get('confidence', 0)}%\n")
                md_parts.append(_html_to_markdown(html_text) if html_text else text)
        content = "\n".join(md_parts)
        filename_encoded = urllib.parse.quote(f"{filename}_ocr.md")
        return Response(
            content=content,
            media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename*=utf-8''{filename_encoded}"},
        )
    elif export_format == "docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise HTTPException(500, "DOCX export requires python-docx") from exc

        doc = Document()
        doc.add_heading(filename, level=1)
        for num in sorted(pages.keys(), key=int):
            p = pages[num]
            text = p.get("text", "")
            html_text = p.get("html_text", "")
            if text or html_text:
                doc.add_heading(
                    f"Trang {num} · {p.get('method', '')} · {p.get('confidence', 0)}%",
                    level=2,
                )
                body = _html_to_markdown(html_text) if html_text else text
                for line in body.splitlines():
                    if line.strip():
                        doc.add_paragraph(line)
                
                # Check for extracted images of this page and embed them into DOCX
                img_dir = UPLOAD_DIR / "extracted_images" / job_id
                if img_dir.exists() and img_dir.is_dir():
                    try:
                        import glob

                        from docx.shared import Inches
                        pattern = str(img_dir / f"page_{num}_img_*")
                        img_files = glob.glob(pattern)
                        
                        def get_img_idx(path_str):
                            import re
                            m = re.search(r'img_(\d+)\.', path_str)
                            return int(m.group(1)) if m else 0
                            
                        img_files.sort(key=get_img_idx)
                        
                        if img_files:
                            doc.add_heading("Hình ảnh trích xuất", level=3)
                            for img_path in img_files:
                                try:
                                    doc.add_picture(img_path, width=Inches(4.5))
                                    doc.add_paragraph(f"Ảnh: {os.path.basename(img_path)}")
                                except Exception as e:
                                    logger.error(f"Error adding picture {img_path} to docx: {e}")
                    except Exception as e:
                        logger.error(f"Failed to process docx image attachment for page {num}: {e}")
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        filename_encoded = urllib.parse.quote(f"{filename}_ocr.docx")
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=utf-8''{filename_encoded}"},
        )
    else:
        lines = []
        for num in sorted(pages.keys(), key=int):
            p = pages[num]
            text = p.get("text", "")
            html_text = p.get("html_text", "")
            body = text or (_html_to_markdown(html_text) if html_text else "")
            if body:
                lines.append(f"--- Page {num} ({p.get('method', '')}, {p.get('confidence', 0)}%) ---")
                lines.append(body)
                lines.append("")
        content = "\n".join(lines)
        filename_encoded = urllib.parse.quote(f"{filename}_ocr.txt")
        return Response(
            content=content,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename*=utf-8''{filename_encoded}"},
        )


def _html_to_markdown(html: str) -> str:
    """Convert HTML to Markdown (headings, bold, italic, tables, lists)."""
    import re
    from html import unescape

    t = html
    # Headings h1–h6
    for i in range(6, 0, -1):
        t = re.sub(
            rf'<h{i}[^>]*>(.*?)</h{i}>',
            lambda m, _i=i: '#' * _i + ' ' + re.sub(r'<[^>]+>', '', m.group(1)).strip(),
            t, flags=re.IGNORECASE | re.DOTALL
        )
    # Bold / italic
    t = re.sub(r'<(?:strong|b)[^>]*>(.*?)</(?:strong|b)>', r'**\1**', t, flags=re.IGNORECASE | re.DOTALL)
    t = re.sub(r'<(?:em|i)[^>]*>(.*?)</(?:em|i)>', r'*\1*', t, flags=re.IGNORECASE | re.DOTALL)
    # Tables → pipe table
    def _table_to_md(m):
        rows = re.findall(r'<tr[^>]*>(.*?)</tr>', m.group(), re.IGNORECASE | re.DOTALL)
        md_rows = []
        for idx, row in enumerate(rows):
            cells = re.findall(r'<(?:td|th)[^>]*>(.*?)</(?:td|th)>', row, re.IGNORECASE | re.DOTALL)
            cells = [re.sub(r'<[^>]+>', '', c).strip() for c in cells]
            md_rows.append('| ' + ' | '.join(cells) + ' |')
            if idx == 0:
                md_rows.append('|' + ' --- |' * len(cells))
        return '\n'.join(md_rows)
    t = re.sub(r'<table[^>]*>.*?</table>', _table_to_md, t, flags=re.IGNORECASE | re.DOTALL)
    # List items
    t = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1', t, flags=re.IGNORECASE | re.DOTALL)
    t = re.sub(r'</?(?:ul|ol)[^>]*>', '', t, flags=re.IGNORECASE)
    # Line breaks and block closings
    t = re.sub(r'<br\s*/?>', '\n', t, flags=re.IGNORECASE)
    t = re.sub(r'</(?:p|div|tr|thead|tbody|h[1-6])>', '\n', t, flags=re.IGNORECASE)
    # Strip all remaining tags
    t = re.sub(r'<[^>]+>', '', t)
    t = unescape(t)
    # Normalise whitespace
    lines = [line.strip() for line in t.split('\n')]
    t = '\n'.join(line for line in lines if line)
    return t.strip()


# ── WebSocket for real-time progress ───────────────────────────────
@app.websocket("/ws/{job_id}")
async def websocket_endpoint(websocket: WebSocket, job_id: str):
    token = websocket.cookies.get("session")
    origin = websocket.headers.get("origin", "")
    if not validate_session(token) or (origin and origin not in CORS_ORIGINS):
        await websocket.close(code=4401)
        return
    job = job_manager.get_job(job_id)
    if not job:
        await websocket.close(code=4004)
        return

    await websocket.accept()
    await job_manager.add_websocket(job_id, websocket)

    # Send current state
    await websocket.send_json({
        "type": "init",
        "job": job.to_dict(),
    })

    try:
        while True:
            # Keep connection alive, wait for client messages
            await websocket.receive_text()
    except WebSocketDisconnect:
        await job_manager.remove_websocket(job_id, websocket)


# ── Simple OCR API (async job pattern, for OpenClaw / external) ────
OCR_API_KEY = os.getenv("OCR_API_KEY", "")

# In-memory store for API OCR jobs (auto-cleanup after 1h)
_api_jobs: dict[str, dict] = {}
_batch_jobs: dict[str, dict] = {}

async def _delayed_cleanup_api_job(job_id: str):
    await asyncio.sleep(3600)
    _api_jobs.pop(job_id, None)

async def _delayed_cleanup_batch_job(batch_id: str):
    await asyncio.sleep(3600)
    _batch_jobs.pop(batch_id, None)

def require_api_key(request: Request):
    """Dependency: require valid API key via X-API-Key header."""
    key = request.headers.get("X-API-Key", "")
    if not OCR_API_KEY or not secrets.compare_digest(key, OCR_API_KEY):
        raise HTTPException(401, "Invalid or missing API key")
    return True


@app.post("/api/v1/ocr")
async def simple_ocr_submit(
    request: Request,
    file: UploadFile = File(...),
    pages: str = Query(default="all", description="Pages: all, 1,3,5, odd, even"),
    method: str = Query(default="auto", description="auto|tesseract|vision"),
    extract_images: bool = Query(default=False, description="Extract original images"),
    _auth: bool = Depends(require_api_key),
):
    """Submit PDF for OCR. Returns job_id for polling."""
    filename = safe_upload_name(file.filename, {".pdf"})
    content = await read_upload_limited(file)

    # Save file
    filepath = UPLOAD_DIR / f"api_{uuid.uuid4().hex[:12]}_{filename}"
    with open(filepath, "wb") as f:
        f.write(content)

    # Analyze PDF
    try:
        analysis = analyze_pdf(str(filepath))
        if analysis.total_pages > MAX_PDF_PAGES:
            raise ValueError(f"PDF has {analysis.total_pages} pages; maximum is {MAX_PDF_PAGES}")
    except Exception as e:
        os.remove(filepath)
        raise HTTPException(500, f"PDF analysis failed: {e}") from e

    # Determine pages
    all_pages = list(range(1, analysis.total_pages + 1))
    if pages == "all":
        selected = all_pages
    elif pages == "odd":
        selected = [p for p in all_pages if p % 2 == 1]
    elif pages == "even":
        selected = [p for p in all_pages if p % 2 == 0]
    else:
        selected = [int(p.strip()) for p in pages.split(",") if p.strip().isdigit()]
        selected = [p for p in selected if 1 <= p <= analysis.total_pages]

    if not selected:
        os.remove(filepath)
        raise HTTPException(400, "No valid pages selected")

    # Create job
    import time as _time
    job_id = uuid.uuid4().hex[:12]
    _api_jobs[job_id] = {
        "status": "processing",
        "filename": filename,
        "filepath": str(filepath),
        "total_pages": len(selected),
        "completed_pages": 0,
        "created_at": _time.time(),
        "html_result": None,
        "error": None,
    }

    # Process in background
    asyncio.create_task(_run_api_ocr_and_cleanup(job_id, str(filepath), filename, selected, method, analysis, extract_images))

    return {
        "job_id": job_id,
        "status": "processing",
        "filename": filename,
        "total_pages": len(selected),
        "poll_url": f"/api/v1/ocr/{job_id}",
    }

async def _run_api_ocr_and_cleanup(job_id, filepath, filename, selected, method, analysis, extract_images):
    async with _job_semaphore:
        await _api_ocr_process(job_id, filepath, filename, selected, method, analysis, extract_images)
    asyncio.create_task(_delayed_cleanup_api_job(job_id))

async def _api_ocr_process(job_id: str, filepath: str, filename: str, selected: list[int], method: str, analysis, extract_images: bool = False):
    """Background: OCR selected pages with 2-phase parallel batch architecture.

    Phase 1: Handle digital/tesseract pages sequentially (fast, no API calls).
    Phase 2: Collect vision-destined pages, batch and run in parallel.
    """
    job = _api_jobs[job_id]
    force = method if method != "auto" else None

    try:
        # Dict to collect results keyed by page_num (for ordered output)
        page_results_map = {}
        vision_pages = []  # collect (page_num, image) for batched vision

        # ── Phase 1: Digital + Tesseract (fast, sequential) ──
        for page_num in selected:
            page_info = next((p for p in analysis.pages if p.page_num == page_num), None)
            classification = page_info.classification if page_info else "scan_complex"

            # Digital pages: extract text directly (no OCR)
            if classification == "digital" and force != "vision":
                text = extract_page_text(filepath, page_num)
                raw_html = f"<pre>{text}</pre>"
                enriched_html = _enrich_html_with_images(job_id, page_num, filepath, raw_html, extract_images)
                
                page_results_map[page_num] = {
                    "page": page_num,
                    "method": "digital",
                    "confidence": 100.0,
                    "text": text,
                    "html_text": enriched_html,
                }
                job["completed_pages"] = len(page_results_map)
                continue

            # Force tesseract: process individually
            if force == "tesseract":
                image = render_page_to_image(filepath, page_num)
                result = await smart_ocr(image, classification, force_method="tesseract")
                
                raw_html = result.get("html_text", f"<pre>{result['text']}</pre>")
                enriched_html = _enrich_html_with_images(job_id, page_num, filepath, raw_html, extract_images)
                
                page_results_map[page_num] = {
                    "page": page_num,
                    "method": result["method"],
                    "confidence": result["confidence"],
                    "text": result.get("text", ""),
                    "html_text": enriched_html,
                }
                job["completed_pages"] = len(page_results_map)
                continue

            # Simple scans in auto mode: try Tesseract first
            if classification == "scan_simple" and force != "vision":
                image = render_page_to_image(filepath, page_num)
                result = await smart_ocr(image, classification, force_method=None)
                if result["method"] == "tesseract":
                    # Tesseract was good enough
                    raw_html = result.get("html_text", f"<pre>{result['text']}</pre>")
                    enriched_html = _enrich_html_with_images(job_id, page_num, filepath, raw_html, extract_images)
                    
                    page_results_map[page_num] = {
                        "page": page_num,
                        "method": result["method"],
                        "confidence": result["confidence"],
                        "text": result.get("text", ""),
                        "html_text": enriched_html,
                    }
                    job["completed_pages"] = len(page_results_map)
                    continue
                # smart_ocr already fell back to vision — use that result
                raw_html = result.get("html_text", f"<pre>{result['text']}</pre>")
                enriched_html = _enrich_html_with_images(job_id, page_num, filepath, raw_html, extract_images)
                
                page_results_map[page_num] = {
                    "page": page_num,
                    "method": result["method"],
                    "confidence": result["confidence"],
                    "text": result.get("text", ""),
                    "html_text": enriched_html,
                }
                job["completed_pages"] = len(page_results_map)
                continue

            # Vision-destined: collect for parallel batching
            image = render_page_to_image(filepath, page_num)
            vision_pages.append((page_num, image))

        # ── Phase 2: Vision batch parallel ──
        if vision_pages:
            batches = [
                vision_pages[i:i + BATCH_SIZE]
                for i in range(0, len(vision_pages), BATCH_SIZE)
            ]

            logger.info(
                f"API Job {job_id}: {len(vision_pages)} vision pages → "
                f"{len(batches)} batches (size={BATCH_SIZE}), "
                f"parallel={PARALLEL_BATCHES}"
            )

            async def _run_batch(batch):
                results = await vision_ocr_batch(batch)
                for (pn, _img), result in zip(batch, results, strict=True):
                    raw_html = result.get("html_text", f"<pre>{result['text']}</pre>")
                    enriched_html = _enrich_html_with_images(job_id, pn, filepath, raw_html, extract_images)
                    
                    page_results_map[pn] = {
                        "page": pn,
                        "method": result["method"],
                        "confidence": result["confidence"],
                        "text": result.get("text", ""),
                        "html_text": enriched_html,
                    }
                    job["completed_pages"] = len(page_results_map)

            # Run batches in parallel waves
            for wave_start in range(0, len(batches), PARALLEL_BATCHES):
                wave = batches[wave_start:wave_start + PARALLEL_BATCHES]
                await asyncio.gather(*[_run_batch(b) for b in wave])

        # ── Build HTML from ordered results ──
        page_results = [page_results_map[p] for p in selected if p in page_results_map]

        html_parts = [
            '<!DOCTYPE html>',
            '<html lang="vi">',
            '<head>',
            f'<title>OCR: {filename}</title>',
            '<meta charset="utf-8">',
            '<style>',
            'body { font-family: Georgia, serif; max-width: 900px; margin: 0 auto; padding: 20px; background: #fafafa; color: #333; }',
            '.page { background: white; padding: 30px; margin: 20px 0; border: 1px solid #ddd; box-shadow: 0 2px 8px rgba(0,0,0,0.08); }',
            '.page-header { border-bottom: 2px solid #eee; padding-bottom: 8px; margin-bottom: 16px; font-size: 14px; color: #888; }',
            'table { border-collapse: collapse; width: 100%; }',
            'td, th { border: 1px solid #ccc; padding: 6px 10px; }',
            '</style>',
            '</head>',
            '<body>',
            f'<h1>📄 {filename}</h1>',
            f'<p style="color:#888">Pages: {len(page_results)} | Method: {method}</p>',
        ]
        for p in page_results:
            html_parts.append('<div class="page">')
            html_parts.append(f'<div class="page-header">Trang {p["page"]} · {p["method"]} · {p["confidence"]}%</div>')
            html_parts.append(p["html_text"])
            html_parts.append('</div>')
        html_parts.append('</body></html>')

        html_content = '\n'.join(html_parts)
        if extract_images:
            html_content = inline_base64_images(html_content, job_id)
        job["html_result"] = html_content
        job["pages"] = {
            str(pn): {
                "text": res.get("text", ""),
                "html_text": res["html_text"],
                "method": res["method"],
                "confidence": res["confidence"]
            }
            for pn, res in page_results_map.items()
        }
        job["status"] = "completed"

    except Exception as e:
        job["status"] = "failed"
        job["error"] = str(e)

    finally:
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
        except OSError:
            pass


@app.get("/api/v1/ocr/{job_id}")
async def simple_ocr_status(
    job_id: str,
    _auth: bool = Depends(require_api_key),
):
    import urllib.parse
    """Poll OCR job status. Returns HTML file when completed."""
    job = _api_jobs.get(job_id)
    if not job:
        raise HTTPException(404, "Job not found or expired")

    if job["status"] == "processing":
        return {
            "job_id": job_id,
            "status": "processing",
            "progress": f"{job['completed_pages']}/{job['total_pages']}",
        }

    if job["status"] == "failed":
        return {
            "job_id": job_id,
            "status": "failed",
            "error": job["error"],
        }

    # Completed — return HTML file
    filename_encoded = urllib.parse.quote(f"{job['filename']}_ocr.html")
    return Response(
        content=job["html_result"],
        media_type="text/html; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename*=utf-8''{filename_encoded}"},
    )


@app.post("/api/v1/ocr/batch")
async def batch_ocr_submit(
    request: Request,
    files: list[UploadFile] = File(...),
    method: str = Query(default="auto", description="auto|tesseract|vision"),
    _auth: bool = Depends(require_api_key),
):
    """Submit multiple PDFs for OCR. Returns batch_id for polling."""
    if not files or len(files) > 10:
        raise HTTPException(400, "Batch must contain between 1 and 10 PDF files")
    batch_id = uuid.uuid4().hex[:12]
    
    saved_files = []
    for f in files:
        filename = safe_upload_name(f.filename, {".pdf"})
        content = await read_upload_limited(f)
        filepath = UPLOAD_DIR / f"batch_{batch_id}_{uuid.uuid4().hex[:12]}_{filename}"
        with open(filepath, "wb") as out:
            out.write(content)
        saved_files.append((str(filepath), filename))
            
    if not saved_files:
        raise HTTPException(400, "No valid PDF files uploaded")
        
    import time as _time
    _batch_jobs[batch_id] = {
        "status": "processing",
        "total_files": len(saved_files),
        "completed_files": 0,
        "results": [],
        "created_at": _time.time(),
        "error": None,
    }
    
    asyncio.create_task(_process_batch_task(batch_id, saved_files, method))
    
    return {
        "batch_id": batch_id,
        "status": "processing",
        "total_files": len(saved_files),
        "poll_url": f"/api/v1/ocr/batch/{batch_id}"
    }

async def _process_batch_task(batch_id: str, saved_files: list, method: str):
    import time as _time
    batch = _batch_jobs[batch_id]
    
    for filepath, filename in saved_files:
        try:
            analysis = analyze_pdf(filepath)
            if analysis.total_pages > MAX_PDF_PAGES:
                raise ValueError(f"PDF has {analysis.total_pages} pages; maximum is {MAX_PDF_PAGES}")
            selected = list(range(1, analysis.total_pages + 1))
            
            job_id = uuid.uuid4().hex[:12]
            _api_jobs[job_id] = {
                "status": "processing",
                "filename": filename,
                "filepath": filepath,
                "total_pages": len(selected),
                "completed_pages": 0,
                "created_at": _time.time(),
                "html_result": None,
                "error": None,
            }
            # Process sequentially to bound memory usage
            async with _job_semaphore:
                await _api_ocr_process(job_id, filepath, filename, selected, method, analysis)
            asyncio.create_task(_delayed_cleanup_api_job(job_id))
            
            job = _api_jobs[job_id]
            batch["results"].append({
                "filename": filename,
                "status": job["status"],
                "html_result": job.get("html_result"),
                "error": job.get("error")
            })
            
        except Exception as e:
            batch["results"].append({
                "filename": filename,
                "status": "failed",
                "error": str(e)
            })
        finally:
            batch["completed_files"] += 1
            if os.path.exists(filepath):
                try:
                    os.remove(filepath)
                except OSError:
                    pass
                    
    batch["status"] = "completed"
    asyncio.create_task(_delayed_cleanup_batch_job(batch_id))


@app.get("/api/v1/ocr/batch/{batch_id}")
async def batch_ocr_status(batch_id: str, _auth: bool = Depends(require_api_key)):
    if batch_id not in _batch_jobs:
        raise HTTPException(404, "Batch job not found")
    data = _batch_jobs[batch_id]
    return {
        "batch_id": batch_id,
        "status": data["status"],
        "total_files": data["total_files"],
        "completed_files": data["completed_files"],
        "error": data["error"]
    }


@app.get("/api/v1/ocr/batch/{batch_id}/download")
async def batch_ocr_download(batch_id: str, format: str = "zip", _auth: bool = Depends(require_api_key)):
    if batch_id not in _batch_jobs:
        raise HTTPException(404, "Batch job not found")
    data = _batch_jobs[batch_id]
    
    if data["status"] != "completed":
        raise HTTPException(400, "Batch job is not yet completed")
        
    if format == "json":
        import json
        return Response(
            content=json.dumps(data["results"], ensure_ascii=False, indent=2),
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="batch_{batch_id}_results.json"'}
        )
        
    # Default to zip containing HTML files
    import io
    import zipfile
    
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        for res in data["results"]:
            fname = res["filename"]
            if res["status"] == "completed" and res["html_result"]:
                zip_file.writestr(f"{fname}.html", res["html_result"].encode('utf-8'))
            else:
                error_msg = res.get("error", "Unknown error")
                zip_file.writestr(f"{fname}_ERROR.txt", error_msg.encode('utf-8'))
                
    return Response(
        content=zip_buffer.getvalue(),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="batch_{batch_id}_results.zip"'}
    )


@app.post("/api/v1/latex/compile")
async def latex_compile_endpoint(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    main_file: str = Form(default="main.tex"),
    engine: str = Form(default="latexmk"),
    timeout: int = Form(default=120),
    _auth: bool = Depends(require_api_key),
):
    """Compile a LaTeX .tex file or .zip project and return the generated PDF."""
    if not LATEX_COMPILE_ENABLED:
        raise HTTPException(503, "LaTeX compilation is disabled")
    filename = safe_upload_name(file.filename or "main.tex", {".tex", ".zip"})
    content = await read_upload_limited(file)

    job_dir = UPLOAD_DIR / f"latex_{uuid.uuid4().hex[:12]}"
    try:
        main_tex = prepare_latex_workspace(job_dir, filename, content, main_file)
        result = await compile_latex_project(
            job_dir=job_dir,
            main_tex=main_tex,
            requested_engine=engine,
            timeout=timeout,
        )
    except LatexCompileError as exc:
        shutil.rmtree(job_dir, ignore_errors=True)
        raise HTTPException(status_code=400, detail=exc.to_detail()) from exc
    except Exception as exc:
        shutil.rmtree(job_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail={"message": str(exc)}) from exc

    background_tasks.add_task(shutil.rmtree, job_dir, True)
    return FileResponse(
        path=result.pdf_path,
        media_type="application/pdf",
        filename=result.download_name,
        headers={
            "X-LaTeX-Engine": result.engine,
        },
        background=background_tasks,
    )


# ── Health Check ────────────────────────────────────────────────────
def collect_health() -> tuple[str, dict]:
    """Collect dependency health without exposing configuration publicly."""
    db_ok = False
    db_error = None
    try:
        from database import DB_PATH
        conn = sqlite3.connect(DB_PATH)
        conn.execute("SELECT 1")
        conn.close()
        db_ok = True
    except Exception as exc:
        db_error = str(exc)

    upload_writable = False
    upload_error = None
    try:
        probe = UPLOAD_DIR / ".healthcheck"
        probe.write_text("ok", encoding="utf-8")
        probe.unlink(missing_ok=True)
        upload_writable = True
    except Exception as exc:
        upload_error = str(exc)

    tesseract_path = shutil.which("tesseract")
    latex_health = get_latex_health()

    checks = {
        "database": {"ok": db_ok, "error": db_error},
        "uploads": {"ok": upload_writable, "error": upload_error},
        "tesseract": {"ok": bool(tesseract_path)},
        "vision": {
            "ok": bool(os.getenv("NINE_ROUTER_URL") and os.getenv("NINE_ROUTER_API_KEY")),
            "model": os.getenv("VISION_MODEL", ""),
        },
        "latex": {
            "ok": not LATEX_COMPILE_ENABLED or latex_health["ok"],
            "enabled": LATEX_COMPILE_ENABLED,
        },
    }
    status = "ok" if all(item["ok"] for item in checks.values()) else "degraded"
    return status, checks


@app.get("/api/health")
async def health():
    status, _checks = collect_health()
    return {
        "status": status,
        "service": "smart-pdf",
        "version": app.version,
    }


@app.get("/api/health/details")
async def health_details(_auth: bool = Depends(require_api_key)):
    status, checks = collect_health()
    return {"status": status, "service": "smart-pdf", "version": app.version, "checks": checks}


# ── Static files (production) ──────────────────────────────────────
DIST_DIR = Path(__file__).resolve().parent.parent / "frontend" / "dist"
if DIST_DIR.exists():
    app.mount("/assets", StaticFiles(directory=DIST_DIR / "assets"), name="static")

    @app.get("/{full_path:path}")
    async def serve_spa(full_path: str):
        """Serve index.html for all non-API routes (SPA fallback)."""
        if full_path == "docs" or full_path == "redoc" or full_path == "openapi.json" or full_path.startswith("api/"):
            raise HTTPException(404, "Not found")
        file_path = (DIST_DIR / full_path).resolve()
        if full_path and file_path.is_relative_to(DIST_DIR.resolve()) and file_path.exists() and file_path.is_file():
            return FileResponse(file_path)
        return FileResponse(DIST_DIR / "index.html")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
